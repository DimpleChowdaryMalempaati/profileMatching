"""Milestone 1 file-system tools for loading and validating resumes.

These functions are designed as LLM-callable tools with consistent return
schemas, explicit error codes, and usage examples in docstrings.

Example (Python):
    >>> from tools.resume_loader import list_resumes, load_resume
    >>> files = list_resumes("resumes")
    >>> doc = load_resume(files["files"][0]["path"])
    >>> print(doc["metadata"]["name"])

Example (CLI-style):
    python -c "from tools.resume_loader import list_resumes; print(list_resumes())"
"""

from __future__ import annotations

import re
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from config import MAX_FILE_SIZE_MB, RESUMES_DIR, SUPPORTED_RESUME_EXTENSIONS


class ResumeLoaderError(Exception):
    """Base exception for resume loading failures."""

    def __init__(self, message: str, code: str = "LOAD_ERROR"):
        super().__init__(message)
        self.code = code


@dataclass
class FileValidationResult:
    valid: bool
    path: str
    extension: str | None = None
    size_bytes: int = 0
    errors: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class ResumeDocument:
    path: str
    filename: str
    format: str
    text: str
    metadata: dict[str, Any]
    sections: dict[str, str] = field(default_factory=dict)
    loaded_at: str = field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def _error_response(
    tool: str, message: str, code: str = "ERROR", **extra: Any
) -> dict[str, Any]:
    """Standard error envelope for all tool outputs."""
    return {
        "success": False,
        "tool": tool,
        "error": {"message": message, "code": code},
        **extra,
    }


def _success_response(tool: str, **payload: Any) -> dict[str, Any]:
    """Standard success envelope for all tool outputs."""
    return {"success": True, "tool": tool, **payload}


def validate_file(path: str | Path) -> dict[str, Any]:
    """Validate a resume file before loading.

    Checks existence, supported extension, and file size limits.

    Args:
        path: Absolute or relative path to a resume file.

    Returns:
        dict with keys: success, tool, valid, path, extension, size_bytes,
        errors, warnings.

    Example:
        >>> result = validate_file("resumes/john_doe.txt")
        >>> if result["valid"]:
        ...     doc = load_resume(result["path"])
    """
    tool = "validate_file"
    file_path = Path(path)

    errors: list[str] = []
    warnings: list[str] = []

    if not file_path.exists():
        return _error_response(
            tool,
            f"File not found: {file_path}",
            code="FILE_NOT_FOUND",
            valid=False,
            path=str(file_path),
        )

    if not file_path.is_file():
        return _error_response(
            tool,
            f"Path is not a file: {file_path}",
            code="NOT_A_FILE",
            valid=False,
            path=str(file_path),
        )

    extension = file_path.suffix.lower()
    size_bytes = file_path.stat().st_size
    size_mb = size_bytes / (1024 * 1024)

    if extension not in SUPPORTED_RESUME_EXTENSIONS:
        errors.append(
            f"Unsupported extension '{extension}'. "
            f"Supported: {sorted(SUPPORTED_RESUME_EXTENSIONS)}"
        )

    if size_mb > MAX_FILE_SIZE_MB:
        errors.append(
            f"File exceeds {MAX_FILE_SIZE_MB} MB limit ({size_mb:.2f} MB)"
        )
    elif size_bytes == 0:
        errors.append("File is empty")

    if extension == ".pdf" and size_bytes < 100:
        warnings.append("PDF file is unusually small; content may be corrupt")

    result = FileValidationResult(
        valid=len(errors) == 0,
        path=str(file_path.resolve()),
        extension=extension,
        size_bytes=size_bytes,
        errors=errors,
        warnings=warnings,
    )
    payload = result.to_dict()
    payload["tool"] = tool
    payload["success"] = result.valid
    if not result.valid:
        payload["error"] = {
            "message": "; ".join(errors),
            "code": "VALIDATION_FAILED",
        }
    return payload


def list_resumes(directory: str | Path | None = None) -> dict[str, Any]:
    """List all supported resume files in a directory.

    Args:
        directory: Folder to scan. Defaults to config RESUMES_DIR.

    Returns:
        dict with success, tool, directory, count, files (list of metadata dicts).

    Example:
        >>> result = list_resumes()
        >>> for f in result["files"]:
        ...     print(f["filename"], f["format"])
    """
    tool = "list_resumes"
    root = Path(directory) if directory else RESUMES_DIR

    if not root.exists():
        return _error_response(
            tool,
            f"Directory not found: {root}",
            code="DIR_NOT_FOUND",
            directory=str(root),
            count=0,
            files=[],
        )

    if not root.is_dir():
        return _error_response(
            tool,
            f"Path is not a directory: {root}",
            code="NOT_A_DIRECTORY",
            directory=str(root),
            count=0,
            files=[],
        )

    files: list[dict[str, Any]] = []
    skipped: list[dict[str, str]] = []

    for file_path in sorted(root.rglob("*")):
        if not file_path.is_file():
            continue
        ext = file_path.suffix.lower()
        if ext not in SUPPORTED_RESUME_EXTENSIONS:
            skipped.append({"path": str(file_path), "reason": "unsupported_extension"})
            continue
        validation = validate_file(file_path)
        entry = {
            "path": str(file_path.resolve()),
            "filename": file_path.name,
            "format": ext.lstrip("."),
            "size_bytes": validation.get("size_bytes", 0),
            "valid": validation.get("valid", False),
        }
        if not entry["valid"]:
            entry["validation_errors"] = validation.get("errors", [])
        files.append(entry)

    return _success_response(
        tool,
        directory=str(root.resolve()),
        count=len(files),
        files=files,
        skipped_count=len(skipped),
        skipped=skipped[:10],
    )


def _extract_text_from_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ResumeLoaderError(
            "pypdf is required for PDF files. Install with: pip install pypdf",
            code="MISSING_DEPENDENCY",
        ) from exc

    try:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
    except Exception as exc:
        raise ResumeLoaderError(
            f"Failed to read PDF: {exc}", code="PDF_READ_ERROR"
        ) from exc

    if not text:
        raise ResumeLoaderError(
            "PDF contains no extractable text (may be scanned/image-only)",
            code="EMPTY_PDF",
        )
    return text


def _extract_text_from_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ResumeLoaderError(
            "python-docx is required for DOCX files. "
            "Install with: pip install python-docx",
            code="MISSING_DEPENDENCY",
        ) from exc

    try:
        document = Document(str(path))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs).strip()
    except Exception as exc:
        raise ResumeLoaderError(
            f"Failed to read DOCX: {exc}", code="DOCX_READ_ERROR"
        ) from exc

    if not text:
        raise ResumeLoaderError("DOCX contains no text", code="EMPTY_DOCX")
    return text


def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".txt", ".md"}:
        try:
            return path.read_text(encoding="utf-8").strip()
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1").strip()
    if ext == ".pdf":
        return _extract_text_from_pdf(path)
    if ext == ".docx":
        return _extract_text_from_docx(path)
    raise ResumeLoaderError(
        f"Unsupported format: {ext}", code="UNSUPPORTED_FORMAT"
    )


SECTION_PATTERNS: dict[str, re.Pattern[str]] = {
    "summary": re.compile(
        r"(?im)^(?:summary|profile|objective|about)\s*:?\s*$"
    ),
    "experience": re.compile(
        r"(?im)^(?:experience|work\s+history|employment|professional\s+experience)\s*:?\s*$"
    ),
    "education": re.compile(
        r"(?im)^(?:education|academic\s+background|qualifications)\s*:?\s*$"
    ),
    "skills": re.compile(
        r"(?im)^(?:skills|technical\s+skills|core\s+competencies|technologies)\s*:?\s*$"
    ),
    "projects": re.compile(
        r"(?im)^(?:projects|personal\s+projects|key\s+projects)\s*:?\s*$"
    ),
    "certifications": re.compile(
        r"(?im)^(?:certifications?|licenses?)\s*:?\s*$"
    ),
}


def _split_into_sections(text: str) -> dict[str, str]:
    """Split resume text into named sections when headers are present."""
    lines = text.splitlines()
    section_starts: list[tuple[int, str]] = []

    for idx, line in enumerate(lines):
        for section_name, pattern in SECTION_PATTERNS.items():
            if pattern.match(line.strip()):
                section_starts.append((idx, section_name))
                break

    if not section_starts:
        return {"full_text": text.strip()}

    sections: dict[str, str] = {}
    for i, (start_idx, name) in enumerate(section_starts):
        end_idx = section_starts[i + 1][0] if i + 1 < len(section_starts) else len(lines)
        body = "\n".join(lines[start_idx + 1 : end_idx]).strip()
        if body:
            sections[name] = body
    return sections


def _extract_name(text: str, filename: str) -> str:
    for line in text.splitlines()[:5]:
        cleaned = line.strip()
        if not cleaned or len(cleaned) > 60:
            continue
        if re.search(r"@|http|linkedin|github|phone|\d{3}[-.\s]\d{3}", cleaned, re.I):
            continue
        if re.match(r"^[A-Z][a-z]+(?: [A-Z][a-z]+)+$", cleaned):
            return cleaned
    stem = Path(filename).stem.replace("_", " ").replace("-", " ")
    return " ".join(word.capitalize() for word in stem.split())


def _extract_skills(text: str, sections: dict[str, str]) -> list[str]:
    skill_block = sections.get("skills", "")
    candidates = re.split(r"[,|•\n;]", skill_block) if skill_block else []
    if not candidates:
        keywords = re.findall(
            r"\b(Python|Java|JavaScript|TypeScript|SQL|AWS|Docker|Kubernetes|"
            r"Machine Learning|Deep Learning|TensorFlow|PyTorch|React|Node\.js|"
            r"Go|Rust|C\+\+|Scala|Spark|Hadoop|PostgreSQL|MongoDB|Redis|"
            r"Azure|GCP|CI/CD|Agile|Scrum|FastAPI|Django|Flask|Spring|"
            r"Data Analysis|Pandas|NumPy|Scikit-learn|NLP|Computer Vision|"
            r"DevOps|Terraform|Ansible|Linux|Git|REST|GraphQL|Microservices|"
            r"Leadership|Communication|Project Management)\b",
            text,
            re.I,
        )
        return sorted({k if k[0].isupper() else k.title() for k in keywords})

    skills = []
    for item in candidates:
        skill = item.strip(" -•\t")
        if skill and 1 < len(skill) < 40:
            skills.append(skill)
    return skills[:30]


def _extract_experience_years(text: str) -> float:
    patterns = [
        r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:experience|exp)",
        r"(?:experience|exp)[:\s]+(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)",
        r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s+(?:in|with)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.I)
        if match:
            return float(match.group(1))

    ranges = re.findall(r"(20\d{2})\s*[-–]\s*(20\d{2}|present|current)", text, re.I)
    if ranges:
        total = 0
        current_year = datetime.now().year
        for start, end in ranges:
            end_year = current_year if end.lower() in {"present", "current"} else int(end)
            total += max(0, end_year - int(start))
        return float(min(total, 40))
    return 0.0


def _extract_education(text: str, sections: dict[str, str]) -> list[str]:
    edu_block = sections.get("education", text)
    degrees = re.findall(
        r"(?:B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|M\.?B\.?A\.?|Ph\.?D\.?|"
        r"Bachelor(?:'s)?|Master(?:'s)?|Doctor(?:ate)?)"
        r"[^,\n]{0,80}(?:Computer Science|Engineering|Mathematics|Business|"
        r"Data Science|Information Technology|Statistics|Economics)?[^,\n]{0,40}",
        edu_block,
        re.I,
    )
    if degrees:
        return [d.strip() for d in degrees[:5]]
    uni_lines = [
        line.strip()
        for line in edu_block.splitlines()
        if re.search(r"university|college|institute|school", line, re.I)
    ]
    return uni_lines[:3]


def extract_metadata(text: str, filename: str, sections: dict[str, str]) -> dict[str, Any]:
    """Extract structured metadata from resume text."""
    return {
        "name": _extract_name(text, filename),
        "skills": _extract_skills(text, sections),
        "experience_years": _extract_experience_years(text),
        "education": _extract_education(text, sections),
        "section_names": list(sections.keys()),
        "word_count": len(text.split()),
    }


def load_resume(path: str | Path) -> dict[str, Any]:
    """Load and parse a single resume file.

    Validates the file, extracts text, sections, and metadata.

    Args:
        path: Path to resume (.txt, .md, .pdf, .docx).

    Returns:
        On success: success=True, document={path, text, metadata, sections, ...}
        On failure: success=False, error={message, code}

    Example:
        >>> doc = load_resume("resumes/alice_smith.txt")
        >>> if doc["success"]:
        ...     print(doc["document"]["metadata"]["name"])
        ...     print(doc["document"]["metadata"]["skills"][:5])
    """
    tool = "load_resume"
    validation = validate_file(path)
    if not validation.get("valid"):
        return _error_response(
            tool,
            validation.get("error", {}).get("message", "Validation failed"),
            code=validation.get("error", {}).get("code", "VALIDATION_FAILED"),
            path=str(path),
        )

    file_path = Path(path)
    try:
        text = _extract_text(file_path)
        sections = _split_into_sections(text)
        metadata = extract_metadata(text, file_path.name, sections)
        document = ResumeDocument(
            path=str(file_path.resolve()),
            filename=file_path.name,
            format=file_path.suffix.lower().lstrip("."),
            text=text,
            metadata=metadata,
            sections=sections,
        )
        return _success_response(tool, document=document.to_dict())
    except ResumeLoaderError as exc:
        return _error_response(tool, str(exc), code=exc.code, path=str(path))
    except Exception as exc:
        return _error_response(
            tool, f"Unexpected error: {exc}", code="UNEXPECTED_ERROR", path=str(path)
        )


def load_all_resumes(directory: str | Path | None = None) -> dict[str, Any]:
    """Load every valid resume in a directory.

    Skips invalid files and collects per-file errors instead of failing entirely.

    Example:
        >>> batch = load_all_resumes("resumes")
        >>> print(batch["loaded_count"], "loaded,", batch["failed_count"], "failed")
    """
    tool = "load_all_resumes"
    listing = list_resumes(directory)
    if not listing.get("success") and listing.get("count", 0) == 0:
        return listing

    documents: list[dict[str, Any]] = []
    failures: list[dict[str, str]] = []

    for file_info in listing.get("files", []):
        if not file_info.get("valid"):
            failures.append(
                {
                    "path": file_info["path"],
                    "error": "; ".join(file_info.get("validation_errors", ["invalid"])),
                }
            )
            continue
        result = load_resume(file_info["path"])
        if result.get("success"):
            documents.append(result["document"])
        else:
            failures.append(
                {
                    "path": file_info["path"],
                    "error": result.get("error", {}).get("message", "unknown"),
                }
            )

    return _success_response(
        tool,
        directory=listing.get("directory"),
        loaded_count=len(documents),
        failed_count=len(failures),
        documents=documents,
        failures=failures,
    )
